"""LLM-free smoke test for the skill core, run in CI.

Exercises the parts that need no API key: multi-source registration, schema
introspection, the SQL + Python sandbox (with the report builder and a chart),
and both report renderers (HTML + Word). Exits non-zero on any failure.
"""

import os
import sys
import tempfile

# Make the package importable when run from the repo root.
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from data_agent import DataRegistry, Sandbox
from data_agent.docx_writer import render_docx
from data_agent.report import render_html


def main() -> None:
    reg = DataRegistry()
    created = reg.register_path("sample-data/sales.csv")
    assert created == ["sales"], created

    schema = reg.schema_summary()
    assert "sales" in schema, schema

    sb = Sandbox(reg)

    sql = sb.run_sql("SELECT count(*) AS n FROM sales")
    assert sql["ok"], sql
    assert "15" in sql["result"], sql["result"]

    py = sb.run_python(
        'g = sql("SELECT region, sum(revenue) AS revenue FROM sales '
        'GROUP BY 1 ORDER BY revenue DESC")\n'
        'report.add_table(g, title="各区域收入")\n'
        'g.plot(kind="bar", x="region", y="revenue", legend=False)\n'
        'report.add_chart(title="各区域收入")\n'
        "len(g)"
    )
    assert py["ok"], py
    assert len(sb.report.blocks) == 2, sb.report.blocks

    narrative = (
        "## 一、执行摘要\n- 测试结论\n\n"
        "## 三、分析框架与口径\n| 项目 | 内容 |\n|---|---|\n| 指标 | 收入 |\n"
    )
    meta = {"question": "q", "model": "ci", "generated_at": "now", "sources": ["sales.csv"]}

    html = render_html(meta, narrative, sb.report.blocks)
    assert "data:image/png;base64," in html, "chart not embedded"
    assert "rpt-table" in html, "table not rendered"
    assert "<h2>一、执行摘要</h2>" in html, "section heading not rendered"

    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "report.docx")
        render_docx(meta, narrative, sb.report.blocks, path)
        from docx import Document

        doc = Document(path)
        assert len(doc.tables) >= 2, f"expected >=2 tables, got {len(doc.tables)}"
        assert len(doc.inline_shapes) >= 1, "expected an embedded chart image"

    reg.close()
    print("CORE SMOKE TEST PASSED")


if __name__ == "__main__":
    main()
