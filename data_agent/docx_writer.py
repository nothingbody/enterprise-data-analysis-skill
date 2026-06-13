"""Render the analysis report as a Word (.docx) document.

Reuses the shared Markdown parser so the Word file mirrors the HTML: the same
structured narrative (with real tables) plus the evidence tables and charts the
agent pinned during analysis.
"""

from __future__ import annotations

import base64
import io
import re
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .report import parse_markdown

_ACCENT = RGBColor(0x25, 0x63, 0xEB)
_MUTED = RGBColor(0x6B, 0x72, 0x80)
_TABLE_STYLE = "Light Grid Accent 1"


def _add_inline(paragraph: Any, text: str) -> None:
    """Add text to a paragraph, honoring **bold** and stripping `code` ticks."""
    text = re.sub(r"`([^`]+?)`", r"\1", text)
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if i % 2 == 1:  # captured group = bold
            run.bold = True


def _add_md_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    ncol = len(headers)
    table = doc.add_table(rows=1, cols=ncol)
    try:
        table.style = _TABLE_STYLE
    except KeyError:
        table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(re.sub(r"[*`]", "", h))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            cells[j].text = re.sub(r"[*`]", "", val)


def _add_nodes(doc: Document, nodes: list[dict]) -> None:
    for n in nodes:
        if n["type"] == "heading":
            doc.add_heading(re.sub(r"[*`]", "", n["text"]), level=max(1, min(n["level"] - 1, 8)))
        elif n["type"] == "list":
            style = "List Number" if n["ordered"] else "List Bullet"
            for item in n["items"]:
                p = doc.add_paragraph(style=style)
                _add_inline(p, item)
        elif n["type"] == "table":
            _add_md_table(doc, n["headers"], n["rows"])
        else:
            p = doc.add_paragraph()
            _add_inline(p, n["text"])


def _add_df(doc: Document, df: pd.DataFrame, max_rows: int) -> None:
    df = df.head(max_rows)
    cols = [str(c) for c in df.columns]
    table = doc.add_table(rows=1, cols=len(cols))
    try:
        table.style = _TABLE_STYLE
    except KeyError:
        table.style = "Table Grid"
    for j, c in enumerate(cols):
        run = table.rows[0].cells[j].paragraphs[0].add_run(c)
        run.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(df.columns):
            val = row[c]
            cells[j].text = "" if pd.isna(val) else str(val)


def _add_chart(doc: Document, png_b64: str) -> None:
    stream = io.BytesIO(base64.b64decode(png_b64))
    doc.add_picture(stream, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_docx(meta: dict[str, Any], narrative: str, blocks: list[dict], path: str) -> str:
    doc = Document()

    title = doc.add_heading("数据分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    q = doc.add_paragraph()
    r = q.add_run("业务问题：")
    r.bold = True
    q.add_run(meta.get("question", ""))

    sources = "、".join(meta.get("sources") or []) or "—"
    meta_p = doc.add_paragraph()
    mr = meta_p.add_run(
        f"数据源：{sources}    生成时间：{meta.get('generated_at', '')}    "
        f"分析模型：{meta.get('model', '')}"
    )
    mr.font.size = Pt(9)
    mr.font.color.rgb = _MUTED

    # Main report body (the structured narrative).
    if narrative.strip():
        _add_nodes(doc, parse_markdown(narrative))

    # Evidence appendix.
    if blocks:
        doc.add_heading("图表与数据明细", level=1)
        for b in blocks:
            if b.get("title"):
                doc.add_heading(b["title"], level=2)
            if b["type"] == "markdown":
                _add_nodes(doc, parse_markdown(b["text"]))
            elif b["type"] == "table":
                _add_df(doc, b["df"], b["max_rows"])
                if b["total_rows"] > b["max_rows"]:
                    note = doc.add_paragraph(
                        f"共 {b['total_rows']} 行，仅展示前 {b['max_rows']} 行。"
                    )
                    note.runs[0].font.size = Pt(9)
                    note.runs[0].font.color.rgb = _MUTED
            elif b["type"] == "chart":
                _add_chart(doc, b["png"])
            if b.get("note"):
                p = doc.add_paragraph()
                _add_inline(p, b["note"])
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = _MUTED

    foot = doc.add_paragraph(
        "本报告由数据分析 Agent 自动生成：基于所提供数据的一次性探索性分析，"
        "图表与数值均来自实际执行的查询。结论供业务参考，正式决策口径请结合指标治理与跨表校验。"
    )
    foot.runs[0].font.size = Pt(8.5)
    foot.runs[0].font.color.rgb = _MUTED

    doc.save(path)
    return path
